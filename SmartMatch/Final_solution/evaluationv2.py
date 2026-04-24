# Evaluation class where firstly i will be doing semantic classification to find the correcttly classified roles returned
# that also represent the same sector as the user cv

from preprocessing import Preprocessing
import pandas as pd
import pdfplumber 
from sklearn.metrics.pairwise import cosine_similarity
from docx import Document
import numpy as np


# imported for visualizing my evaluations
import matplotlib

# TkAgg is a Matplotlib backend which renders my plots using 
# a GUI window can be interacted with
matplotlib.use('TkAgg')

import matplotlib.pyplot as plt

import pandas as pd
# displays dataframe information at full / not cut off (using for testing)
pd.set_option('display.max_colwidth', None)


class Evaluation:

    # constructor 
    def __init__(self):
        self.preprocessor = Preprocessing()
        self.preprocessor.readcsv("SmartMatch/Data/postings.csv")
        self.preprocessor.combine_relevant_fields(['title', 'location', 'skills_desc', 'description'])
        self.resume_df = None

        # Exploratory evaluation framework (industry-to-category semantic mapping using embedding similarity as proxy labels)
        # self.job_industries_df = None


    def get_recommended_roles(self,k_best):
        
        # returns all rows from df by their index from k_best.
        recommended_jobs = self.preprocessor.df.iloc[k_best]

        # Stores the job titles of all recommended roles
        recommended_titles = recommended_jobs["title"].values

        # Stores the job ids of all recommended roles
        recommended_job_ids = recommended_jobs["job_id"].values

        return recommended_job_ids, recommended_titles
    
    
    # Exploratory evaluation framework (industry-to-category semantic mapping using embedding similarity as proxy labels)
    def jobid_industry_dataset(self, file_loc):

        # Stores job_industries.csv as a dataframe
        self.job_industries_df = pd.read_csv(file_loc)

        # Stores industries.csv as a dataframe
        industries = pd.read_csv("SmartMatch/Data/industries.csv")

        # drop missing values
        self.job_industries_df.dropna(inplace= True)

        # drop missing values from industries dataset
        industries.dropna(inplace = True)

        # convert values to lower case
        industries["industry_name"] =  industries["industry_name"].str.lower()

        # merge datasets on the industry id.
        job_industry_dataset = pd.merge(self.job_industries_df, industries, on='industry_id') 

        ###   filter job_ids so only the ones in pf remain   ###

        # Stores a set of all job_ids in pf
        valid_jobs = set(self.preprocessor.pf['job_id'])

        # Mask to check which job id in job_industry_dataset is in dataset pf
        mask = job_industry_dataset["job_id"].isin(valid_jobs)
        
        # Apply mask to job_industry dataset so relevant job_ids remain
        job_industry_dataset = job_industry_dataset[mask]

        return job_industry_dataset
    

     # Implemented to load the resume CV into a dataframe, clean the dataset etc.
    

    def cv_dataset(self, file_loc):

        # store dataset as df
        self.resume_df = pd.read_csv(file_loc)

        # removed unimportant attribute from the data frame, leaving
        # only relevant fields 
        del self.resume_df['Resume_html']

        # convert all attribute values into lower case for consistency with job postings.
        self.resume_df["Resume_str"] = self.resume_df["Resume_str"].str.lower()
        self.resume_df["Category"] = self.resume_df["Category"].str.lower()

        # Drop the rows where at least one element is missing.
        self.resume_df.dropna(inplace= True)

        # To avoid class imbalance, grouped observations by their categories, and from each group, select 20 observations at random
        # NOTE: random_state = 42 (controlled randomness)
        new_df = self.resume_df.groupby("Category").sample(n = 20, random_state= 42)

        # resets the index into the new dataframe so it starts from 0.
        new_df = new_df.reset_index(drop=True)

        return new_df
    
    # Works out the overlap@k of roles per cv between both model
    def overlap_at_k(self,tfidf, embeddings,k):

        # convert each list to a set to only store unique job ids
        tfidf_set = set(tfidf)
        embeddings_set = set(embeddings)

        # Store the intersection of job_ids between the models 
        intersection = tfidf_set.intersection(embeddings_set)

        # Measures the proportion of shared items in the top-K results
        return len(intersection) / k
    
    # Works out the how similar the recommendation sets are, relative to 
    # all unique roles across both models
    def jaccard_similarity(self,tfidf, embeddings):

        # convert each list to a set to only store unique job ids
        tfidf_set = set(tfidf)
        embeddings_set = set(embeddings)

        # Store the intersection of job_ids between the models
        intersection = tfidf_set.intersection(embeddings_set)

        # Store the union of job_ids between the models
        union = tfidf_set.union(embeddings_set)

        return len(intersection) / len(union)

    # Evaluating both TF-IDF and Sentence Embeddings models
    # in a single function using: Top-K Overlap (Overlap@K), Jaccard Similarity

    def evaluate_models(self, cv_data, k):

        # each index contains results per inputted CV
        overlap_results = []

        # each index contains results per inputted CV
        jaccard_results = []

        # Vectorize job postings dataset (TF-IDF)
        X_tfidf = self.preprocessor.convert_postings_TFIDF()

        # Vectorize job postings dataset (Sentence embeddings)
        X_embeddings = self.preprocessor.convert_postings_embeddings()

        # Instantiate embeddings model
        model = self.preprocessor.get_model()

        # Vectorize CV dataset for TF-IDF and Embeddings
        for i in range(len(cv_data)):

            # get the current CV observation
            current_row = cv_data.iloc[i]

            # Retrieve CV text
            cv_text = current_row["Resume_str"]

            # Using TF-IDF vectorize CV text
            Y_tfidf = self.preprocessor.convert_user_TFIDF(cv_text)

            # Using embeddings vectorize CV text
            # reshape the structure to match job postings embeddings
            Y_embeddings = model.encode(cv_text).reshape(1, -1)

            # apply cosine similarity to get vector values between
            # each model and the job postings
            tfidf_similarity = cosine_similarity(X_tfidf, Y_tfidf)
            embeddings_similarity = cosine_similarity(X_embeddings, Y_embeddings)

            tfidf_scores = tfidf_similarity.flatten()
            embeddings_scores = embeddings_similarity.flatten()

            kbest_tfidf = np.argsort(tfidf_scores)[-k:][::-1]

            kbest_embeddings = np.argsort(embeddings_scores)[-k:][::-1]

            # Retrieve job_ids and job titles for top k results 
            tfidf_ids, tfidf_titles = self.get_recommended_roles(kbest_tfidf)
            embeddings_ids, embeddings_titles = self.get_recommended_roles(kbest_embeddings)

            # evaluation #

            # Top K Overlap
            k_overlap = self.overlap_at_k(tfidf_ids, embeddings_ids, k)

            # Jaccard Similarity
            jaccard_s = self.jaccard_similarity(tfidf_ids, embeddings_ids)

            # Add results to list of results per CV input
            overlap_results.append(k_overlap)
            jaccard_results.append(jaccard_s)

        # mean overlap across all CVs
        mean_overlap = sum(overlap_results) / len(overlap_results)
        mean_jaccard = sum(jaccard_results) / len(jaccard_results)

        return mean_overlap, mean_jaccard
    
    def mean_jaccard_graph(self, k_values, jaccard_means):

        # Plot line graph to show the change in jaccard similarity between models
        # as K increases 

        plt.figure()
        plt.plot(k_values, jaccard_means, marker='o')
        plt.xlabel('K')
        plt.ylabel('Mean Jaccard Similarity')
        plt.title('Jaccard Similarity vs K')
        plt.xticks(k_values)
        plt.ylim(0, max(jaccard_means) + 0.01)
        plt.grid()

        plt.savefig("SmartMatch/Graphs/final_evaluation/jaccard_vs_k.png")
        plt.show()

    def mean_overlap_graph(self, k_values, overlap_means):

        # Plot line graph to show the change in overlap between models
        # as K increases 

        plt.figure()
        plt.plot(k_values, overlap_means, marker='o')
        plt.xlabel('K')
        plt.ylabel('Mean Overlap@K')
        plt.title('Overlap@K vs K')
        plt.xticks(k_values)
        plt.ylim(0, max(overlap_means) + 0.01)
        plt.grid()

        plt.savefig("SmartMatch/Graphs/final_evaluation/overlap_vs_k.png")
        plt.show()




    # Exploratory evaluation framework (proxy-based sector labelling using manually defined mappings)
    def cv_sector_dict(self):
        cv_and_sectors = {
            'CV1.docx': 'Tech',
            'CV2.docx': 'Tech',
            'CV3.docx': 'Tech',
            'CV4.docx': 'Engineering',
            'CV5.docx': 'Finance',
            'CV6.docx': 'Finance',
            'CV7.docx': 'Law',
            'CV8.docx': 'Engineering',
            'CV9.docx': 'Operations',
            'CV10.docx': 'Finance'
            }
        return cv_and_sectors
    
    # Exploratory evaluation framework (proxy-based sector labelling using manually defined mappings)
    def sectors_desc(self):
        sectors = {
            'Law': "Legal jobs including attorneys, paralegals, compliance and corporate law roles.",
            'Finance': "Finance jobs including accounting, auditing, banking, investment analysis, financial planning.",
            'Tech':"Technology jobs including software engineering, data science, machine learning, IT support, web development, AI engineering.",
            'Operations': "Operations jobs including project management, logistics, coordination and business operations.",
            'Engineering': "Engineering jobs including mechanical, civil, electrical, manufacturing and industrial engineering."
            }
        return list(sectors.keys()), list(sectors.values())

    # Exploratory evaluation framework (industry-to-category semantic mapping using embedding similarity as proxy labels)
    def run_tfidf2(self, cleaned_dataset, industries_dataset):

        # store the categories in a sorted list where each index is unique
        categories = sorted(list(set(cleaned_dataset["Category"])))

        # The model implemented to embed job industries and categories to derive the 
        # predicted category for each job
        # Retrieved the model from preprocessing for consistency
        model = self.preprocessor.get_model()

        # encoded the list of categories which will be used to work out similarity score 
        # between each job industry and all encoded categories.
        encoded_categories = model.encode(categories)

        # the number of roles returned to each user
        k = 5

        # stores precisoon of K roles recommended per CV
        # where each index represents the result for a different CV.
        all_precisions = []

        # vectorize the job postings dataset using TF-IDF
        X = self.preprocessor.convert_postings_TFIDF()

        # For loop that goes through each CV in the dataset
        for i in range(len(cleaned_dataset)):
            count = 0

            # get the current CV observation
            current_row = cleaned_dataset.iloc[i]

            # Retrieve CV text
            cv_text = current_row["Resume_str"]

            # Retrieve CV Category (Ground Truth)
            ground_truth = current_row["Category"]

            # Using TF-IDF vectorize CV text
            Y = self.preprocessor.convert_user_TFIDF(cv_text)

            # Compute similarity score betwwen job postings and CV vectors 
            similarity_score = cosine_similarity(X,Y)

            # flatten the 2-D array to get a series of similarity scores fpr sorting.
            scores = similarity_score.flatten() 

            # returns 'k' indexs of the roles with the highest similarity scores within the dataset. 
            # argsort sorts the scores in ascending order, keeping the index which it was initially stored in.
            k_best = np.argsort(scores)[-k:][::-1]

            # get the job_id and title from recommended roles
            job_ids, titles = self.get_recommended_roles(k_best)

            ## sector classification begins ##

            # Retrieve the industries for each role by the job_ids. 
            # This will be used for the classification

            # Returns a list of true or false checking which job ids in the
            # industries dataset are also in k job_ids
            mask = industries_dataset["job_id"].isin(job_ids)

            # applies the mask to dataset, only returning observations that 
            # returned true
            k_industries = industries_dataset[mask]

            # Due to one-to-many relationship between job_ids and industries,
            # group by job_id and store all industries for that job in a corresponding list
            grouped = k_industries.groupby("job_id")["industry_name"].apply(list)

            # For loop that vectorizes industries per job_id
            for job_id, industries_list in grouped.items():

                # Vectorize the industries using sentence embeddings
                encoded_industries = model.encode(industries_list)

                # get the similarity scores between the industry and encoded categories
                score = cosine_similarity(encoded_industries, encoded_categories)

                # Some jobs have multiple industries so first get the indexes
                # of the highest matches per industry
                # NOTE: axis = 1 as where comparing across the industry itself
                best_per_industry = score.argmax(axis = 1)

                # Get the actual max score value for each industry
                best_scores = score.max(axis = 1)

                # Get the best industry overall 
                best_industry_index = best_scores.argmax()

                # Get the index of the category from the now chosen
                # best industry
                category_index = best_per_industry[best_industry_index]

                # Get the predicted category 
                predicted_category = categories[category_index]

                # Check predicted category against current CVs category (ground truth)
                if predicted_category == ground_truth:
                    count += 1

                # print results
                title = self.preprocessor.df.loc[self.preprocessor.df["job_id"] == job_id, "title"].values[0]
                print(f"For CV: {i}, ROLE {title}, sector classified: {predicted_category}")
                print("GT:", ground_truth, "| Pred:", predicted_category)
                
            # Get the precision of correctly classified jobs out of all
            # recommended jobs for that CV
            precision = (count / len(grouped))
            all_precisions.append(precision)

            # return the precisison
            print(f"The precision of correctly indentified roles within the same sector as the User CV: {precision}")
            print("\n")
        
        # return the average precision across all the cvs. 
        average_precision = sum(all_precisions) / len(all_precisions)
        print(f"The Overall precision:{average_precision * 100}%")

    # Exploratory evaluation framework (proxy-based sector labelling using manually defined mappings)
    def run_tfidf(self):

        # stores the sector names and values in seperate lists.
        # sector_names will be used to predict/classify the sector to the role
        # sector_values will be used for semantic embeddings and compared with each role recommended to the user. 
        sector_names, sector_values = self.sectors_desc()
        
        # The model implemeted to embed the sector desc for the classification. 
        # Retrieved the model from preprocessing for consistency.
        model = self.preprocessor.get_model()


        # the list is encoded which will be used to work out the similarity score for
        # each role against all sectors
        encoded_sector_desc = model.encode(sector_values)

        # stores precision of K roles recommended per CV where 
        # each index represents the results for a different CV.
        all_precisions = []

        # the number of roles returned to each user.
        k = 5
        
        # vectorize job postings
        X = self.preprocessor.convert_postings_TFIDF()

        # stores cv dictionary in cv_dict for processing.
        cv_dict = self.cv_sector_dict()

        for key in cv_dict.keys():
            count = 0
            text = ""

            if key.endswith(".pdf"):
                with pdfplumber.open(f"SmartMatch/Sample_cvs/{key}") as pdf:
               # loop to extract the text from each page within the pdf
               # if the page is empty, appends "" to the string 'text'.
                 for page in pdf.pages:
                    text += page.extract_text() or ""
                    text = text.lower()

            elif key.endswith(".docx"):
                document = Document(f"SmartMatch/Sample_cvs/{key}")
                for p in document.paragraphs:
                    text += p.text
                    text = text.lower()
            
            # vectorize user cv (TF-IDF)
            Y = self.preprocessor.convert_user_TFIDF(text)

            # Compute similarity score betwwen job postings and CV
            similarity_score = cosine_similarity(X,Y)

            # Flatten the similarity_score array 2-D -> 1-D to allow for easier sorting. 
            scores = similarity_score.flatten() 

            # returns 'k' indexs of the roles with the highest similarity scores within the dataset. 
            # argsort sorts the scores in ascending order, keeping the index which it was initially stored in. 
            k_best = np.argsort(scores)[-k:][::-1]

            # gets top k roles
            # store recommended titles in a list
            recommended_titles = self.get_recommended_roles(k_best)

            # sector classification begins # 
            # this is done for each role recommended per cv against each embedded sector
            for title in recommended_titles:

                # NOTE: stores the encoded job titles in a new variable
                # used [] around the input so it output is 2-D. as originally just a string.
                encoded_titles = model.encode([title])

                # get the similarity scores between the job and the sector desc
                score = cosine_similarity(encoded_titles, encoded_sector_desc)

                # flatten converts 2-D array -> 1-D array
                match_scores = score.flatten()

                # gets the index of the highest score, which is the derived sector for that role.
                label = np.argmax(match_scores)
                predicted_sector = sector_names[label]

                # Check if the predicted sector for that role is the same as the sector
                # belonging to the cv 
                if predicted_sector == cv_dict[key]:
                    count += 1
                print(f"For CV {key}, ROLE {title}, sector classified: {predicted_sector}")

            # Get the precision of ranked roles being within the same sector as the user cv
            precision = (count / k)
            all_precisions.append(precision)

            # return the precisison
            print(f"The precision of correctly indentified roles within the same sector as the User CV: {precision}")
            print("\n")

        # return the average precision across all the cvs. 
        average_precision = sum(all_precisions) / len(all_precisions)
        print(f"The Overall precision:{average_precision * 100}%")

        return all_precisions, (average_precision)
    
    # Exploratory evaluation framework (proxy-based sector labelling using manually defined mappings)
    def run_embeddings(self):

        # stores the sector names and values in seperate lists.
        # sector_names will be used to predict/classify the sector to the role
        # sector_values will be used for semantic embeddings and compared with each role recommended to the user. 
        sector_names, sector_values = self.sectors_desc()

        # The model implemeted to embed the sector desc for the classification, job postings and cvs.
        # retrieved the model from pre-processing.
        model = self.preprocessor.get_model()

        # the list is encoded which will be used to work out the similarity score for
        # each role against all sectors
        encoded_sector_desc = model.encode(sector_values)

        # stores precision of K roles recommended per CV where 
        # each index represents the results for a different CV.
        all_precisions = []

        # the number of roles returned to each user.
        k = 5

        # embed job postings
        X = self.preprocessor.convert_postings_embeddings()

        # stores cv dictionary in cv_dict for processing.
        cv_dict = self.cv_sector_dict()

        for key in cv_dict.keys():
            count = 0
            text = ""

            if key.endswith(".pdf"):
                with pdfplumber.open(f"SmartMatch/Sample_cvs/{key}") as pdf:
               # loop to extract the text from each page within the pdf
               # if the page is empty, appends "" to the string 'text'.
                 for page in pdf.pages:
                    text += page.extract_text() or ""
                    text = text.lower()

            elif key.endswith(".docx"):
                document = Document(f"SmartMatch/Sample_cvs/{key}")
                for p in document.paragraphs:
                    text += p.text
                    text = text.lower()
            
            # vectorize user cv (embedding)
            Y = self.preprocessor.convert_user_embeddings(text)

            # Compute similarity score betwwen job postings and CV
            similarity_score = cosine_similarity(X,Y)

            # Flatten the similarity_score array 2-D -> 1-D to allow for easier sorting. 
            scores = similarity_score.flatten() 

            # returns 'k' indexs of the roles with the highest similarity scores within the dataset. 
            # argsort sorts the scores in ascending order, maining the index which it was stored in. 
            k_best = np.argsort(scores)[-k:][::-1]

            # gets top k roles
            # store recommended titles in a list
            recommended_titles = self.get_recommended_roles(k_best)

            # sector classification begins # 
            # this is done for each role recommended per cv against each embedded sector
            for title in recommended_titles:

                # NOTE: stores the encoded job titles in a new variable
                # used [] around the input so it output is 2-D. as originally just a string.
                encoded_titles = model.encode([title])

                # get the similarity scores between the job and the sector desc
                score = cosine_similarity(encoded_titles, encoded_sector_desc)

                # flatten converts 2-D array -> 1-D array
                match_scores = score.flatten()

                # gets the index of the highest score, which is the derived sector for that role.
                label = np.argmax(match_scores)
                predicted_sector = sector_names[label]

                # Check if the predicted sector for that role is the same as the sector
                # belonging to the cv 
                if predicted_sector == cv_dict[key]:
                    count += 1
                print(f"For CV {key}, ROLE {title}, sector classified: {predicted_sector}")
            
            # Get the precision of ranked roles being within the same sector as the user cv
            precision = (count / k)
            all_precisions.append(precision)

            # return the precisison
            print(f"The precision of correctly indentified roles within the same sector as the User CV: {precision}")
            print("\n")

        # return the average precision across all the cvs. 
        average_precision = sum(all_precisions) / len(all_precisions)
        print(f"The Overall precision:{average_precision * 100}%")

        return all_precisions, (average_precision)

    # Visualisation of exploratory evaluation results (not part of final model comparison)
    def precision_graph_models(self, tfidf, embeddings):

        # retrieves cvs and their corresponding sectors.
        cv_and_sectors = self.cv_sector_dict()
       
        # removes "_cv.pdf and or "_cv.docx" for each cv in a dictionary ror better X labels
        plot_cv = [i.replace('_cv.pdf', '').replace('_cv.docx', '') for i in cv_and_sectors.keys()]


        # anomylising the user cvs for ethics reasons
        new_plots = [plot_cv[i].replace(plot_cv[i],f"CV {i+1}") for i in range(len(plot_cv))]

        # list of the precisions per CVs' (TF-IDF)
        results_tfidf = tfidf

        # list of the precisions per CVs' (Sentence Embeddings)
        results_embeddings = embeddings

        x = np.arange(len(new_plots))
        w = 0.35

        plt.bar(x - w/2, results_tfidf, w, label = 'TF-IDF')
        plt.bar(x + w/2, results_embeddings, w, label = 'all-MiniLM-L6-v2')
        plt.xticks(x,new_plots)
        plt.xlabel('CVs')
        plt.ylabel('Precision (0-1)')
        plt.title('Precision of Sector Classification for Top-K recommendations')
        plt.legend()

        # prevents labels being clipped when saving the image
        plt.savefig("SmartMatch/Graphs/legacy_evaluation/precision_per_cv.png", dpi = 'figure')
        plt.show()
    
    # Visualisation of exploratory evaluation results (not part of final model comparison)
    def mean_precision_graph(self, mean_tfidf, mean_embeddings):

        # retrieves cvs and their corresponding sectors.
        # cv_and_sectors = self.cv_sector_dict()

        fig, ax = plt.subplots()

        models = ['TF-IDF' , 'all-MiniLM-L6-v2']
        results = [mean_tfidf, mean_embeddings]
        bar_labels = ['TF-IDF ', 'all-MiniLM-L6-v2']
        bar_colors = ['tab:red', 'tab:blue']

        ax.bar(models, results, label=bar_labels, color=bar_colors)


        ax.set_xlabel('Models')
        ax.set_ylabel('Mean Precision (0-1)')
        ax.set_title('Mean Precision per CV For Each Model')
        
        ax.legend()


        # prevents labels being clipped when saving the image
        plt.savefig("SmartMatch/Graphs/legacy_evaluation/average_precision_per_model.png", dpi = 'figure')
        plt.show()

    


# testing
eval = Evaluation()

# industry cv
lookup_table = eval.jobid_industry_dataset("SmartMatch/Data/job_industries.csv")

# cv dataset
cv_data = eval.cv_dataset("SmartMatch/Data/Resume.csv")

k_values = [5,10,20,50,100]

jaccard_means = []
overlap_means = []
results = []

# loop through k_values and compute mean jaccard and overlap for each iteration of K
for k in k_values:
    print(f"For k = {k} (START)")
    jaccard, overlap = eval.evaluate_models(cv_data, k)

    # stores Mean jaccard and overlap at K in list
    # Used for creating graph representations
    jaccard_means.append(jaccard)
    overlap_means.append(overlap)
    
    # Stores the results per iteration to
    # be represented in a dataframe
    results.append(
    {"K": k,
     "Jaccard": jaccard,
     "Overlap": overlap
     }
     )
    print(f'For k = {k} (DONE)')


# store results in a datafrome
results_df = pd.DataFrame(results)
results_df.to_csv("SmartMatch/Data/results.csv", index= False)

# plot results on line graph
eval.mean_overlap_graph(k_values, overlap_means)
eval.mean_jaccard_graph(k_values, jaccard_means)














